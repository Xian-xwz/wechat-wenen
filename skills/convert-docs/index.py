#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档转换为 Markdown 工具
支持 PDF、DOCX、DOC 格式转换
"""

import os
import sys
from pathlib import Path

# ============ PDF 转换 ============
def convert_pdf_to_md(pdf_path, output_path=None):
    """使用 pdfplumber 将 PDF 转换为 Markdown"""
    import pdfplumber
    
    if output_path is None:
        output_path = pdf_path.rsplit('.', 1)[0] + '.md'
    
    print(f"[PDF] 正在转换: {pdf_path}")
    
    with pdfplumber.open(pdf_path) as pdf:
        all_content = []
        for i, page in enumerate(pdf.pages):
            print(f"  处理第 {i+1}/{len(pdf.pages)} 页...")
            text = page.extract_text()
            if text and text.strip():
                all_content.append(f"## 第 {i+1} 页\n\n{text}")
            
            # 尝试提取表格
            tables = page.extract_tables()
            for table in tables:
                if table:
                    all_content.append("\n" + table_to_markdown(table))
    
    content = "\n\n".join(all_content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"[OK] PDF 转换完成: {output_path}\n")
    return output_path

def table_to_markdown(table):
    """将表格转换为 Markdown 格式"""
    if not table or not table[0]:
        return ""
    
    def clean_cell(cell):
        if cell is None:
            return ""
        return str(cell).replace("\n", " ").replace("|", "\\|").strip()
    
    lines = []
    header = [clean_cell(c) for c in table[0]]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + "|".join(["---"] * len(header)) + "|")
    
    for row in table[1:]:
        cells = [clean_cell(c) for c in row]
        while len(cells) < len(header):
            cells.append("")
        lines.append("| " + " | ".join(cells[:len(header)]) + " |")
    
    return "\n".join(lines)

# ============ DOCX 转换 ============
def convert_docx_to_md(docx_path, output_path=None):
    """使用 python-docx 将 DOCX 转换为 Markdown"""
    from docx import Document
    
    if output_path is None:
        output_path = docx_path.rsplit('.', 1)[0] + '.md'
    
    print(f"[DOCX] 正在转换: {docx_path}")
    
    doc = Document(docx_path)
    
    md_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # 根据样式判断标题级别
        style_name = para.style.name if para.style else ""
        
        if 'Heading 1' in style_name or '标题 1' in style_name:
            md_content.append(f"# {text}")
        elif 'Heading 2' in style_name or '标题 2' in style_name:
            md_content.append(f"## {text}")
        elif 'Heading 3' in style_name or '标题 3' in style_name:
            md_content.append(f"### {text}")
        elif text.startswith('第') and ('章' in text or '节' in text):
            md_content.append(f"## {text}")
        else:
            md_content.append(text)
    
    # 处理表格
    for table in doc.tables:
        md_content.append("")
        md_content.append(convert_docx_table_to_md(table))
        md_content.append("")
    
    content = "\n\n".join(md_content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"[OK] DOCX 转换完成: {output_path}\n")
    return output_path

def convert_docx_table_to_md(table):
    """将 DOCX 表格转换为 Markdown"""
    if not table.rows:
        return ""
    
    lines = []
    rows = list(table.rows)
    
    # 表头
    header_cells = [cell.text.replace("\n", " ").replace("|", "\\|").strip() for cell in rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("|" + "|".join(["---"] * len(header_cells)) + "|")
    
    # 数据行
    for row in rows[1:]:
        cells = [cell.text.replace("\n", " ").replace("|", "\\|").strip() for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(lines)

# ============ DOC 转换 (旧版 Word) ============
def convert_doc_to_md(doc_path, output_path=None):
    """转换旧版 .doc 文件"""
    import subprocess
    import tempfile
    
    if output_path is None:
        output_path = doc_path.rsplit('.', 1)[0] + '.md'
    
    print(f"[DOC] 正在转换: {doc_path}")
    
    # 尝试使用 LibreOffice 转换为 docx
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'docx', '--outdir', tmpdir, doc_path],
                capture_output=True,
                text=True
            )
            
            if result.returncode == 0:
                docx_path = os.path.join(tmpdir, os.path.basename(doc_path).replace('.doc', '.docx'))
                if os.path.exists(docx_path):
                    convert_docx_to_md(docx_path, output_path)
                    print(f"[OK] DOC (通过 LibreOffice) 转换完成: {output_path}\n")
                    return output_path
    except FileNotFoundError:
        pass
    
    # 如果没有 LibreOffice，尝试直接读取
    print("  ⚠️ 未找到 LibreOffice，尝试直接读取...")
    
    try:
        with open(doc_path, 'rb') as f:
            content = f.read()
            # 尝试提取文本
            text = content.decode('gbk', errors='ignore')
            # 清理无用字符
            lines = [line.strip() for line in text.split('\n') if line.strip() and len(line.strip()) > 2]
            clean_text = '\n\n'.join(lines[:500])  # 限制行数避免垃圾数据
            
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# {os.path.basename(doc_path)}\n\n")
            f.write(clean_text)
        
        print(f"[OK] DOC (原始读取) 转换完成: {output_path}")
        print(f"   ⚠️ 警告：可能丢失格式，建议用 LibreOffice 重新转换\n")
        return output_path
    except Exception as e:
        print(f"[ERROR] 转换失败: {e}")
        return None

# ============ 主程序 ============
def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='文档转 Markdown 工具')
    parser.add_argument('files', nargs='+', help='要转换的文件路径')
    parser.add_argument('-o', '--output', help='输出目录')
    
    args = parser.parse_args()
    
    print("=" * 50)
    print("文档转 Markdown 工具")
    print("=" * 50 + "\n")
    
    for file in args.files:
        if not os.path.exists(file):
            print(f"[ERROR] 文件不存在: {file}")
            continue
        
        # 确定输出路径
        if args.output:
            output_path = os.path.join(args.output, os.path.basename(file).rsplit('.', 1)[0] + '.md')
        else:
            output_path = None
        
        if file.endswith('.pdf'):
            convert_pdf_to_md(file, output_path)
        elif file.endswith('.docx'):
            convert_docx_to_md(file, output_path)
        elif file.endswith('.doc'):
            convert_doc_to_md(file, output_path)
        else:
            print(f"[ERROR] 不支持的文件格式: {file}")
    
    print("=" * 50)
    print("所有文件转换完成！")
    print("=" * 50)

if __name__ == "__main__":
    main()
